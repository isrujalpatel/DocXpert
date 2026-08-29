"""
text_replacer.py — Find & replace engine (feature 3).

Searches across paragraphs, tables, and headers/footers.
Preserves formatting by replacing text within individual runs.
"""

import re
from pathlib import Path
from dataclasses import dataclass

from utils.helpers import generate_uuid
from config.settings import settings


@dataclass
class ReplaceResult:
    """Result of a find & replace operation."""
    output_path: Path
    replacements_made: int


def _replace_in_runs(runs, find_text: str, replace_text: str,
                     use_regex: bool, case_sensitive: bool) -> int:
    """
    Replace text within a list of runs, preserving formatting.

    Strategy:
    1. Try per-run replacement first (handles matches within a single run).
    2. If the match spans multiple runs, coalesce text, find match,
       then distribute replacement back across runs.

    Returns the number of replacements made.
    """
    total = 0
    flags = 0 if case_sensitive else re.IGNORECASE

    # Phase 1: Per-run replacement (preserves formatting perfectly)
    for run in runs:
        original_text = run.text
        if not original_text:
            continue

        if use_regex:
            new_text, count = re.subn(find_text, replace_text, original_text, flags=flags)
        else:
            if case_sensitive:
                count = original_text.count(find_text)
                new_text = original_text.replace(find_text, replace_text)
            else:
                pattern = re.escape(find_text)
                new_text, count = re.subn(pattern, replace_text, original_text, flags=re.IGNORECASE)

        if count > 0:
            run.text = new_text
            total += count

    # Phase 2: Cross-run replacement
    # Only if we haven't found anything yet and the match might span runs
    if total == 0 and runs and not use_regex:
        full_text = "".join(run.text for run in runs)

        if case_sensitive:
            match_count = full_text.count(find_text)
        else:
            match_count = len(re.findall(re.escape(find_text), full_text, re.IGNORECASE))

        if match_count > 0:
            # Find match positions in the full text
            if case_sensitive:
                positions = [m.start() for m in re.finditer(re.escape(find_text), full_text)]
            else:
                positions = [m.start() for m in re.finditer(re.escape(find_text), full_text, re.IGNORECASE)]

            # Apply replacements in reverse order to preserve positions
            for pos in reversed(positions):
                _replace_across_runs(runs, pos, len(find_text), replace_text)
                total += 1

    return total


def _replace_across_runs(runs, start_pos: int, length: int, replacement: str):
    """
    Replace text that spans multiple runs.
    Inserts the replacement into the first overlapping run and
    removes matched characters from subsequent runs.
    """
    current_pos = 0
    started = False
    remaining = length

    for run in runs:
        run_len = len(run.text)

        if not started:
            if current_pos + run_len > start_pos:
                offset_in_run = start_pos - current_pos
                chars_in_this_run = min(run_len - offset_in_run, remaining)

                run.text = (
                    run.text[:offset_in_run]
                    + replacement
                    + run.text[offset_in_run + chars_in_this_run:]
                )

                remaining -= chars_in_this_run
                started = True

                if remaining <= 0:
                    return
        else:
            chars_to_remove = min(len(run.text), remaining)
            run.text = run.text[chars_to_remove:]
            remaining -= chars_to_remove

            if remaining <= 0:
                return

        current_pos += run_len


def replace_in_docx(
    file_path: Path,
    find_text: str,
    replace_text: str,
    use_regex: bool = False,
    case_sensitive: bool = True,
) -> ReplaceResult:
    """
    Find and replace text in a DOCX document.

    Searches across:
    - Body paragraphs
    - Table cells
    - Headers and footers

    Preserves all formatting (bold, italic, font, color, etc.)
    by replacing text within individual runs.
    """
    from docx import Document

    doc = Document(str(file_path))
    total_replacements = 0

    # 1. Body paragraphs
    for para in doc.paragraphs:
        if para.runs:
            total_replacements += _replace_in_runs(
                para.runs, find_text, replace_text, use_regex, case_sensitive
            )

    # 2. Tables — iterate all tables → rows → cells → paragraphs → runs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.runs:
                        total_replacements += _replace_in_runs(
                            para.runs, find_text, replace_text, use_regex, case_sensitive
                        )

    # 3. Headers and footers — iterate all sections
    for section in doc.sections:
        # Headers
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    if para.runs:
                        total_replacements += _replace_in_runs(
                            para.runs, find_text, replace_text, use_regex, case_sensitive
                        )

        # Footers
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    if para.runs:
                        total_replacements += _replace_in_runs(
                            para.runs, find_text, replace_text, use_regex, case_sensitive
                        )

    # Save modified document
    output_name = f"{generate_uuid()}.docx"
    output_path = settings.UPLOAD_DIR / output_name
    doc.save(str(output_path))

    return ReplaceResult(
        output_path=output_path,
        replacements_made=total_replacements,
    )


def replace_in_document(
    file_path: Path,
    find_text: str,
    replace_text: str,
    use_regex: bool = False,
    case_sensitive: bool = True,
) -> ReplaceResult:
    """
    Find and replace text in a document (dispatches by file type).
    """
    ext = file_path.suffix.lower().lstrip(".")

    if ext == "docx":
        return replace_in_docx(file_path, find_text, replace_text, use_regex, case_sensitive)
    elif ext == "pdf":
        raise ValueError("Find & replace is not supported for PDF files. Convert to DOCX first.")
    else:
        raise ValueError(f"Find & replace is not supported for .{ext} files.")
