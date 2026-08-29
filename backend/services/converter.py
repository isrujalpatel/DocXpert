"""
converter.py — Format conversion service (DOCX ↔ PDF).

DOCX→PDF: Uses LibreOffice headless (best fidelity), falls back to reportlab.
PDF→DOCX: Uses pdf2docx (best fidelity), falls back to PyMuPDF text extraction.
"""

import subprocess
import shutil
from pathlib import Path

from utils.helpers import generate_uuid
from config.settings import settings


def _run_libreoffice(input_path: Path, target_format: str) -> Path:
    """
    Run LibreOffice headless to convert a document.

    Args:
        input_path: Path to the source file.
        target_format: Target format (e.g., 'pdf', 'docx').

    Returns:
        Path to the converted file in the uploads directory.

    Raises:
        RuntimeError if LibreOffice is not available or conversion fails.
    """
    soffice_path = settings.get_libreoffice_path()
    if not soffice_path:
        raise RuntimeError("LibreOffice is not installed or not found on this system.")

    settings.ensure_dirs()

    # LibreOffice outputs to a directory, not a specific file
    temp_dir = settings.TEMP_DIR / generate_uuid()
    temp_dir.mkdir(parents=True, exist_ok=True)

    try:
        cmd = [
            soffice_path,
            "--headless",
            "--norestore",
            "--convert-to", target_format,
            "--outdir", str(temp_dir),
            str(input_path),
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60,  # 60 second timeout
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed: {result.stderr or result.stdout}"
            )

        # Find the output file
        expected_ext = f".{target_format}"
        output_files = list(temp_dir.glob(f"*{expected_ext}"))

        if not output_files:
            raise RuntimeError("LibreOffice produced no output file.")

        # Move to uploads directory with a UUID name
        output_name = f"{generate_uuid()}{expected_ext}"
        output_path = settings.UPLOAD_DIR / output_name
        shutil.move(str(output_files[0]), str(output_path))

        return output_path

    finally:
        # Clean up temp directory
        shutil.rmtree(str(temp_dir), ignore_errors=True)


def convert_docx_to_pdf(input_path: Path) -> Path:
    """
    Convert a DOCX file to PDF.

    Strategy:
    1. Try LibreOffice headless (best fidelity)
    2. Fall back to reportlab (text-only, lower fidelity)
    """
    # Try LibreOffice first
    try:
        return _run_libreoffice(input_path, "pdf")
    except RuntimeError:
        pass  # Fall through to reportlab

    # Fallback: reportlab (basic text extraction)
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    output_name = f"{generate_uuid()}.pdf"
    output_path = settings.UPLOAD_DIR / output_name

    doc = Document(str(input_path))

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y_position = height - 50
    line_height = 14

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y_position -= line_height
            continue

        # Determine font size from paragraph style
        font_size = 11
        font_name = "Helvetica"

        if para.runs and para.runs[0].font.size:
            font_size = min(para.runs[0].font.size.pt, 36)

        if para.style and "Heading" in (para.style.name or ""):
            font_size = max(font_size, 16)
            font_name = "Helvetica-Bold"

        c.setFont(font_name, font_size)

        # Simple word wrapping
        words = text.split()
        line = ""
        for word in words:
            test_line = f"{line} {word}".strip()
            if c.stringWidth(test_line, font_name, font_size) < width - 100:
                line = test_line
            else:
                c.drawString(50, y_position, line)
                y_position -= line_height + 2
                line = word

                if y_position < 50:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y_position = height - 50

        if line:
            c.drawString(50, y_position, line)
            y_position -= line_height + 4

        if y_position < 50:
            c.showPage()
            y_position = height - 50

    c.save()
    return output_path


def convert_pdf_to_docx(input_path: Path) -> Path:
    """
    Convert a PDF file to DOCX.

    Strategy:
    1. Try pdf2docx (best fidelity — preserves layout, tables, images)
    2. Fall back to PyMuPDF text extraction
    """
    output_name = f"{generate_uuid()}.docx"
    output_path = settings.UPLOAD_DIR / output_name

    # Try pdf2docx first
    try:
        from pdf2docx import Converter

        cv = Converter(str(input_path))
        cv.convert(str(output_path))
        cv.close()

        if output_path.exists() and output_path.stat().st_size > 0:
            return output_path
    except Exception:
        pass  # Fall through to PyMuPDF

    # Fallback: PyMuPDF text extraction
    try:
        import fitz  # PyMuPDF

        pdf = fitz.open(str(input_path))

        from docx import Document
        doc = Document()

        for page in pdf:
            text = page.get_text("text")
            if text:
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)

        pdf.close()
        doc.save(str(output_path))
        return output_path

    except Exception:
        pass

    # Last resort: PyPDF2
    from PyPDF2 import PdfReader
    from docx import Document

    reader = PdfReader(str(input_path))
    doc = Document()

    for page in reader.pages:
        text = page.extract_text()
        if text:
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

    doc.save(str(output_path))
    return output_path


def convert_document(input_path: Path, target_format: str) -> Path:
    """
    Convert a document to the specified target format.

    Args:
        input_path: Path to the source document.
        target_format: Target format ('docx' or 'pdf').

    Returns:
        Path to the converted file.

    Raises:
        ValueError: If the conversion is not supported.
    """
    source_ext = input_path.suffix.lower().lstrip(".")

    if source_ext == "docx" and target_format == "pdf":
        return convert_docx_to_pdf(input_path)
    elif source_ext == "pdf" and target_format == "docx":
        return convert_pdf_to_docx(input_path)
    elif source_ext == target_format:
        raise ValueError(f"Source and target format are the same: .{target_format}")
    else:
        raise ValueError(
            f"Conversion from .{source_ext} to .{target_format} is not supported."
        )
