"""
document_parser.py — DOC/DOCX/PDF parsing logic.

Reads uploaded document files and extracts structured content
including paragraphs, fonts, styles, tables, and metadata.
"""

from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field


@dataclass
class ParsedParagraph:
    """A single paragraph extracted from a document."""
    text: str
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    is_bold: bool = False
    is_italic: bool = False
    style_name: Optional[str] = None


@dataclass
class ParsedDocument:
    """Structured representation of a parsed document."""
    paragraphs: List[ParsedParagraph] = field(default_factory=list)
    fonts_used: List[dict] = field(default_factory=list)
    page_count: int = 0
    word_count: int = 0
    metadata: dict = field(default_factory=dict)


def parse_docx(file_path: Path) -> ParsedDocument:
    """
    Parse a DOCX file and extract structured content.

    Args:
        file_path: Path to the .docx file.

    Returns:
        ParsedDocument with paragraphs, fonts, and metadata.
    """
    from docx import Document

    doc = Document(str(file_path))
    parsed = ParsedDocument()
    fonts_seen = {}

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        # Extract font info from the first run
        font_name = None
        font_size = None
        is_bold = False
        is_italic = False

        if para.runs:
            run = para.runs[0]
            if run.font.name:
                font_name = run.font.name
            if run.font.size:
                font_size = run.font.size.pt
            is_bold = run.bold or False
            is_italic = run.italic or False

        parsed.paragraphs.append(ParsedParagraph(
            text=para.text,
            font_name=font_name,
            font_size=font_size,
            is_bold=is_bold,
            is_italic=is_italic,
            style_name=para.style.name if para.style else None,
        ))

        # Track fonts
        if font_name:
            key = f"{font_name}_{font_size}"
            if key not in fonts_seen:
                fonts_seen[key] = {"font_name": font_name, "font_size": font_size, "occurrences": 0}
            fonts_seen[key]["occurrences"] += 1

    parsed.fonts_used = list(fonts_seen.values())
    parsed.word_count = sum(len(p.text.split()) for p in parsed.paragraphs)
    parsed.metadata = {
        "core_properties": {
            "author": doc.core_properties.author or "",
            "title": doc.core_properties.title or "",
        }
    }

    return parsed


def parse_pdf(file_path: Path) -> ParsedDocument:
    """
    Parse a PDF file and extract text content.
    Uses PyMuPDF (fitz) for better text extraction quality.

    Args:
        file_path: Path to the .pdf file.

    Returns:
        ParsedDocument with paragraphs and metadata.
    """
    parsed = ParsedDocument()

    try:
        import fitz  # PyMuPDF

        pdf = fitz.open(str(file_path))
        parsed.page_count = len(pdf)

        for page in pdf:
            text = page.get_text("text")
            if text:
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        parsed.paragraphs.append(ParsedParagraph(text=line))

        parsed.metadata = {
            "page_count": parsed.page_count,
            "pdf_info": pdf.metadata or {},
        }
        pdf.close()

    except ImportError:
        # Fallback to PyPDF2
        from PyPDF2 import PdfReader

        reader = PdfReader(str(file_path))
        parsed.page_count = len(reader.pages)

        for page in reader.pages:
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        parsed.paragraphs.append(ParsedParagraph(text=line))

        parsed.metadata = {
            "page_count": parsed.page_count,
            "pdf_info": reader.metadata._data if reader.metadata else {},
        }

    parsed.word_count = sum(len(p.text.split()) for p in parsed.paragraphs)
    return parsed


def parse_document(file_path: Path) -> ParsedDocument:
    """
    Parse a document file based on its extension.

    Args:
        file_path: Path to the document file.

    Returns:
        ParsedDocument with extracted content.

    Raises:
        ValueError: If the file type is not supported.
    """
    ext = file_path.suffix.lower().lstrip(".")

    if ext == "docx":
        return parse_docx(file_path)
    elif ext == "pdf":
        return parse_pdf(file_path)
    elif ext == "doc":
        # .doc (legacy Word) — would need python-docx or libreoffice conversion
        raise ValueError(
            "Legacy .doc format requires conversion. "
            "Please convert to .docx first or use the conversion endpoint."
        )
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
