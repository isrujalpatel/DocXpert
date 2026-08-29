"""
change_applier.py — Apply accepted AI suggestions to a DOCX document.

Feature 8: Takes a list of accepted SuggestedChange items and applies
them to the document, preserving formatting where possible.
"""

import re
from pathlib import Path
from typing import List

from config.settings import settings
from models.document import SuggestedChange, ChangeLocation
from utils.helpers import generate_uuid


def apply_changes(file_path: Path, changes: List[SuggestedChange]) -> Path:
    """
    Apply accepted changes to a DOCX document.

    Changes are sorted in reverse order by paragraph_index and char_offset
    to prevent position shifts from invalidating subsequent change locations.

    Args:
        file_path: Path to the source .docx file.
        changes: List of accepted SuggestedChange objects.

    Returns:
        Path to the new document with changes applied.
    """
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document(str(file_path))

    # Sort changes in reverse order to apply from end to start
    sorted_changes = sorted(
        changes,
        key=lambda c: (c.location.paragraph_index, c.location.char_offset),
        reverse=True,
    )

    for change in sorted_changes:
        try:
            if change.type in ("spelling", "grammar"):
                _apply_text_change(doc, change)
            elif change.type == "font":
                _apply_font_change(doc, change)
            elif change.type == "margin":
                _apply_margin_change(doc, change)
            elif change.type in ("heading", "formatting"):
                _apply_formatting_change(doc, change)
            elif change.type == "spacing":
                _apply_spacing_change(doc, change)
            elif change.type in ("layout", "list_style"):
                _apply_formatting_change(doc, change)
        except (IndexError, AttributeError, ValueError):
            # Skip changes that can't be applied (stale positions, etc.)
            continue

    # Save as new file
    output_name = f"{generate_uuid()}.docx"
    output_path = settings.UPLOAD_DIR / output_name
    doc.save(str(output_path))

    return output_path


def _apply_text_change(doc, change: SuggestedChange):
    """Apply a text-level change (spelling, grammar) to a specific paragraph."""
    from docx import Document

    para_idx = change.location.paragraph_index
    if para_idx < 0 or para_idx >= len(doc.paragraphs):
        return

    para = doc.paragraphs[para_idx]

    # Try to find and replace within individual runs to preserve formatting
    original = change.original
    suggested = change.suggested

    for run in para.runs:
        if original in run.text:
            run.text = run.text.replace(original, suggested, 1)
            return

    # Fallback: reconstruct paragraph text if original spans multiple runs
    full_text = para.text
    if original in full_text:
        # Find which runs contain the text
        pos = full_text.find(original)
        if pos >= 0:
            _replace_across_runs(para.runs, pos, len(original), suggested)


def _replace_across_runs(runs, start_pos: int, length: int, replacement: str):
    """
    Replace text that spans across multiple runs.
    Modifies the first run containing the start of the match,
    clears text from subsequent runs that overlap.
    """
    current_pos = 0
    started = False
    remaining = length

    for i, run in enumerate(runs):
        run_len = len(run.text)

        if not started:
            if current_pos + run_len > start_pos:
                # This run contains the start of the match
                offset_in_run = start_pos - current_pos
                chars_in_this_run = min(run_len - offset_in_run, remaining)

                run.text = (
                    run.text[:offset_in_run]
                    + replacement
                    + run.text[offset_in_run + chars_in_this_run:]
                )

                remaining -= chars_in_this_run
                started = True

                if remaining <= 0:
                    return
        else:
            # Subsequent runs — remove matched characters
            chars_to_remove = min(run_len, remaining)
            run.text = run.text[chars_to_remove:]
            remaining -= chars_to_remove

            if remaining <= 0:
                return

        current_pos += run_len


def _apply_font_change(doc, change: SuggestedChange):
    """Apply a font change (name, size) to the document."""
    from docx.shared import Pt

    suggested = change.suggested.lower()

    # Parse font name and size from the suggestion
    # e.g., "Times New Roman, 12pt" or "Change font to Arial 11pt"
    import re
    size_match = re.search(r'(\d+(?:\.\d+)?)\s*pt', suggested)
    new_size = float(size_match.group(1)) if size_match else None

    # Extract font name (everything before the size or comma)
    font_name = None
    name_patterns = [
        r'(?:font[:\s]+)?([A-Z][a-zA-Z\s]+?)(?:\s*,|\s*\d)',
        r'"([^"]+)"',
        r"'([^']+)'",
    ]
    for pattern in name_patterns:
        match = re.search(pattern, change.suggested)
        if match:
            font_name = match.group(1).strip()
            break

    para_idx = change.location.paragraph_index

    if para_idx >= 0 and para_idx < len(doc.paragraphs):
        # Apply to specific paragraph
        para = doc.paragraphs[para_idx]
        for run in para.runs:
            if font_name:
                run.font.name = font_name
            if new_size:
                run.font.size = Pt(new_size)
    else:
        # Apply document-wide
        for para in doc.paragraphs:
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if new_size:
                    run.font.size = Pt(new_size)


def _apply_margin_change(doc, change: SuggestedChange):
    """Apply margin changes to document sections."""
    from docx.shared import Inches
    import re

    suggested = change.suggested.lower()

    # Parse margin values: "top: 1in, bottom: 1in, left: 1.25in, right: 1.25in"
    margins = {}
    for side in ("top", "bottom", "left", "right"):
        match = re.search(rf'{side}[:\s]+(\d+(?:\.\d+)?)\s*(?:in|inch)', suggested)
        if match:
            margins[side] = float(match.group(1))

    if not margins:
        return

    for section in doc.sections:
        if "top" in margins:
            section.top_margin = Inches(margins["top"])
        if "bottom" in margins:
            section.bottom_margin = Inches(margins["bottom"])
        if "left" in margins:
            section.left_margin = Inches(margins["left"])
        if "right" in margins:
            section.right_margin = Inches(margins["right"])


def _apply_formatting_change(doc, change: SuggestedChange):
    """Apply a formatting/heading change to a paragraph."""
    para_idx = change.location.paragraph_index

    if para_idx < 0 or para_idx >= len(doc.paragraphs):
        return

    para = doc.paragraphs[para_idx]
    suggested = change.suggested

    # Try to apply as a style name
    style_names = ["Heading 1", "Heading 2", "Heading 3", "Heading 4",
                   "Normal", "Title", "Subtitle", "List Bullet", "List Number"]

    for style in style_names:
        if style.lower() in suggested.lower():
            try:
                para.style = doc.styles[style]
            except KeyError:
                pass
            return


def _apply_spacing_change(doc, change: SuggestedChange):
    """Apply spacing changes to a paragraph."""
    from docx.shared import Pt
    import re

    para_idx = change.location.paragraph_index
    suggested = change.suggested.lower()

    # Parse line spacing
    line_match = re.search(r'line\s*spacing[:\s]+(\d+(?:\.\d+)?)', suggested)
    before_match = re.search(r'before[:\s]+(\d+(?:\.\d+)?)\s*pt', suggested)
    after_match = re.search(r'after[:\s]+(\d+(?:\.\d+)?)\s*pt', suggested)

    if para_idx >= 0 and para_idx < len(doc.paragraphs):
        paras = [doc.paragraphs[para_idx]]
    else:
        paras = doc.paragraphs

    for para in paras:
        fmt = para.paragraph_format
        if line_match:
            fmt.line_spacing = float(line_match.group(1))
        if before_match:
            fmt.space_before = Pt(float(before_match.group(1)))
        if after_match:
            fmt.space_after = Pt(float(after_match.group(1)))
